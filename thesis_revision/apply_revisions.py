# -*- coding: utf-8 -*-
"""Apply revised prose to the thesis, preserving paragraph style, alignment, and
run font. Original file is never modified. Writes a revision log CSV."""
import copy, csv
from docx import Document
from docx.shared import Pt
from revisions_litreview import REV
from revisions_litreview2 import REV2
from revisions_methodology import REVM

ALL = {}
ALL.update(REV); ALL.update(REV2); ALL.update(REVM)

SRC = "Thesis_TARGET.docx"
OUT = "Nawar_TARDOUST_Academically_Humanised.docx"

doc = Document(SRC)
paras = doc.paragraphs
log = []

for i, new_text in ALL.items():
    p = paras[i]
    orig = p.text
    r0 = p.runs[0] if p.runs else None
    font_name = (r0.font.name if r0 else None) or "Times New Roman"
    font_size = r0.font.size if (r0 and r0.font.size) else None
    # remove all existing runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # add single new run
    nr = p.add_run(new_text)
    nr.font.name = font_name
    # set rFonts for ascii/hAnsi/cs so it renders consistently
    from docx.oxml.ns import qn
    rpr = nr._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii','w:hAnsi','w:cs'):
        rfonts.set(qn(a), font_name)
    if font_size:
        nr.font.size = font_size
    log.append({"paragraph_index": i, "original_text": orig, "revised_text": new_text})

doc.save(OUT)

with open("Nawar_TARDOUST_Revision_Log.csv", "w", newline="", encoding="utf-8") as f:
    w = csv.DictWriter(f, fieldnames=["paragraph_index","original_text","revised_text"])
    w.writeheader()
    for row in log:
        w.writerow(row)

print(f"Applied {len(ALL)} paragraph revisions.")
print(f"Saved: {OUT}")
print(f"Log: Nawar_TARDOUST_Revision_Log.csv ({len(log)} rows)")
