# -*- coding: utf-8 -*-
"""Insert full Chapters 5 and 6 into the thesis, replacing the 'opening bridge'
stub, immediately before the References heading. Match Chapter 4 formatting
exactly (bold Times New Roman 16/14/12pt headings; justified 12pt body).
Never modify the original file."""
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from chapters_content import CHAPTERS

SRC = "VF_CH_1_3_4_ORIGINAL.docx"
OUT = "Nawar_TARDOUST_Complete_Thesis_Chapters_1_to_6.docx"

doc = Document(SRC)
body = doc.element.body

# ---- locate the References heading (Heading 1, text 'References') as anchor ----
refs_p = None
for p in doc.paragraphs:
    if p.style.name == "Heading 1" and p.text.strip() == "References":
        refs_p = p
        break
assert refs_p is not None, "References heading not found"
refs_el = refs_p._p

# ---- remove the 'opening bridge' stub: from the stub title up to (excluding) References ----
removed = []
in_stub = False
for p in list(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("Chapter 5 — Discussion (opening bridge)"):
        in_stub = True
    if in_stub:
        if p._p is refs_el:
            break
        removed.append(t[:60])
        p._p.getparent().remove(p._p)
print(f"Removed {len(removed)} stub paragraphs:")
for r in removed:
    print("   -", repr(r))

def set_tnr(run, bold=False, size=None):
    run.font.name = "Times New Roman"
    # ensure complex/east-asian too
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), "Times New Roman")
    run.bold = bold
    if size:
        run.font.size = Pt(size)

def make_para(kind, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if kind == "title":
        pf.page_break_before = True          # page break before each chapter
        pf.space_before = Pt(6); pf.space_after = Pt(12)
        r = p.add_run(text); set_tnr(r, bold=True, size=16)
    elif kind == "h2":
        pf.space_before = Pt(12); pf.space_after = Pt(6)
        r = p.add_run(text); set_tnr(r, bold=True, size=14)
    elif kind == "h3":
        pf.space_before = Pt(8); pf.space_after = Pt(4)
        r = p.add_run(text); set_tnr(r, bold=True, size=12)
    else:  # body
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(6)
        r = p.add_run(text); set_tnr(r, bold=False)
    return p._p

# ---- build new paragraphs (appended to end) then move before References ----
new_els = [make_para(kind, text) for kind, text in CHAPTERS]
for el in new_els:
    el.getparent().remove(el)          # detach from end
    refs_el.addprevious(el)            # insert before References, in order

# ---- ensure a page break precedes the References heading too (clean separation) ----
# (leave References as-is; it already starts a section in the source layout)

# ---- set updateFields=true so Word offers to refresh the TOC on open ----
# updateFields must be placed in schema order: after characterSpacingControl,
# before footnotePr/endnotePr/compat/rsids.
settings = doc.settings.element
if settings.find(qn('w:updateFields')) is None:
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
    anchor = None
    for tag in ('w:footnotePr', 'w:endnotePr', 'w:compat', 'w:rsids', 'w:themeFontLang'):
        el = settings.find(qn(tag))
        if el is not None:
            anchor = el; break
    if anchor is not None:
        anchor.addprevious(uf)
    else:
        settings.append(uf)

doc.save(OUT)
print(f"\nSaved: {OUT}")
print(f"Total paragraphs now: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
